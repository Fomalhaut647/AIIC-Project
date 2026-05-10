"""file_parse tests — Spec E §7.3."""
from pathlib import Path

import pytest

from services.file_parse import parse_file


FIXTURES = Path(__file__).parent / "fixtures"


@pytest.mark.asyncio
async def test_parse_md():
    text, warnings = await parse_file(FIXTURES / "sample.md", "md")
    assert "项目动机" in text
    assert "baseline" in text
    assert warnings == []


@pytest.mark.asyncio
async def test_parse_txt():
    text, warnings = await parse_file(FIXTURES / "sample.txt", "txt")
    assert "纯文本测试" in text
    assert warnings == []


@pytest.mark.asyncio
async def test_parse_pdf_simple(tmp_path: Path):
    """构造一个简单 PDF（内含一行文本），验证解析。"""
    import fitz
    pdf_path = tmp_path / "simple.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "Hello PDF 世界", fontsize=12)
    doc.save(str(pdf_path))
    doc.close()

    text, warnings = await parse_file(pdf_path, "pdf")
    assert "Hello PDF" in text or "世界" in text
    assert warnings == []  # 简单文本，无图片


@pytest.mark.asyncio
async def test_parse_pdf_with_image(tmp_path: Path):
    """含图片的 PDF 应在 warnings 中提示。"""
    import fitz
    pdf_path = tmp_path / "with_image.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "page text", fontsize=12)
    # 插入一个最小像素图（红色方块）
    pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 10, 10))
    pix.set_rect(pix.irect, (255, 0, 0))
    rect = fitz.Rect(100, 100, 110, 110)
    page.insert_image(rect, pixmap=pix)
    doc.save(str(pdf_path))
    doc.close()

    text, warnings = await parse_file(pdf_path, "pdf")
    assert "page text" in text  # 图片处理不应破坏文本提取
    assert any("image" in w.lower() for w in warnings)


@pytest.mark.asyncio
async def test_parse_pdf_encrypted_raises(tmp_path: Path):
    """加密 PDF 抛 ValueError。"""
    import fitz
    pdf_path = tmp_path / "enc.pdf"
    doc = fitz.open()
    doc.new_page().insert_text((50, 50), "secret")
    doc.save(str(pdf_path), encryption=fitz.PDF_ENCRYPT_AES_256, owner_pw="x", user_pw="y")
    doc.close()

    with pytest.raises(ValueError, match="encrypted"):
        await parse_file(pdf_path, "pdf")


@pytest.mark.asyncio
async def test_parse_docx(tmp_path: Path):
    """构造一个简单 docx，验证段落 + 表格。"""
    from docx import Document
    docx_path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("第一段：项目动机")
    doc.add_paragraph("第二段：baseline 选择")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "指标"
    table.rows[0].cells[1].text = "数值"
    table.rows[1].cells[0].text = "准确率"
    table.rows[1].cells[1].text = "0.85"
    doc.save(str(docx_path))

    text, warnings = await parse_file(docx_path, "docx")
    assert "项目动机" in text
    assert "baseline 选择" in text
    assert "指标" in text  # 表格 row 0 被渲染
    assert "准确率" in text  # 表格 row 1 被渲染
    assert any("table" in w.lower() for w in warnings)


@pytest.mark.asyncio
async def test_parse_unsupported_type_raises():
    with pytest.raises(ValueError, match="unsupported"):
        await parse_file(FIXTURES / "sample.txt", "xyz")
