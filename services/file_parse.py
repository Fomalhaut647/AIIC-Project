"""File parsing for project material uploads — Plan3 G1 (Spec E §7.3)."""
from __future__ import annotations

import asyncio
import zipfile
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document  # python-docx
from docx.opc.exceptions import OpcError  # python-docx 解析错误基类


async def parse_file(path: Path, file_type: str) -> tuple[str, list[str]]:
    """根据 file_type 分发；返回 (parsed_text, warnings)。

    file_type 必须是 'pdf' | 'docx' | 'md' | 'txt'。
    PDF 加密 / 解析错误 → ValueError。

    PyMuPDF / python-docx 是 sync 库，10MB 文件最坏阻塞 1-3s；用
    asyncio.to_thread offload 到线程池避免阻塞 fastapi event loop。
    md/txt 直读小开销，不上线程。
    """
    if file_type == "pdf":
        return await asyncio.to_thread(_parse_pdf, path)
    if file_type == "docx":
        return await asyncio.to_thread(_parse_docx, path)
    if file_type in ("md", "txt"):
        return path.read_text(encoding="utf-8"), []
    raise ValueError(f"unsupported file_type: {file_type}")


def _parse_pdf(path: Path) -> tuple[str, list[str]]:
    """PyMuPDF 抽页面文本；图片不 OCR，warnings 提示。"""
    warnings: list[str] = []
    chunks: list[str] = []
    with fitz.open(path) as doc:
        if doc.is_encrypted:
            raise ValueError("PDF is encrypted; please remove password protection")
        for i, page in enumerate(doc):
            txt = page.get_text("text")
            chunks.append(txt)
            if page.get_images():
                warnings.append(f"page {i + 1} contains images (OCR not performed)")
    return "\n\n".join(chunks).strip(), warnings


def _parse_docx(path: Path) -> tuple[str, list[str]]:
    """python-docx 抽段落 + 表格。"""
    warnings: list[str] = []
    # python-docx 解析损坏 .docx 至少抛三类异常（都不继承 ValueError）：
    #   1. PackageNotFoundError (OpcError 子类) — phys_pkg 层未识别 zip
    #   2. zipfile.BadZipFile — zip 结构损坏 / 非 zip 字节
    #   3. KeyError — valid zip 但缺 [Content_Types].xml（用户 mv archive.zip
    #      resume.docx 场景；从 zipfile.getinfo() 直接冒出未 wrap）
    # endpoint 层 catch (ValueError, RuntimeError) narrow except 不覆盖以上
    # 任一 → 逃成 fastapi 500 + raw 文件 orphan leak quota（spec §4 raw+json
    # 配对契约 broken）。service 层 wrap 成 ValueError 让 endpoint 走 422 +
    # unlink raw 路径，与 _parse_pdf 的 "encrypted" ValueError 模式对齐。
    try:
        doc = Document(str(path))
    except (OpcError, zipfile.BadZipFile, KeyError) as e:
        raise ValueError("docx file is corrupted or not a valid .docx") from e
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    if doc.tables:
        for t in doc.tables:
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells]
                parts.append(" | ".join(cells))
        warnings.append("docx contains tables (rendered as plain text rows)")
    return "\n\n".join(parts).strip(), warnings
